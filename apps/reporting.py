from collections import defaultdict
from datetime import date, timedelta
from io import BytesIO

from django.db.models import Count, Q, Sum
from django.utils import timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Project, Report, Task
from .task_visibility import visible_tasks_for


REPORT_TEMPLATES = (
    {
        "type": Report.Type.WEEKLY_PROGRESS,
        "title": "Weekly Progress",
        "description": "Summary of completed and in-progress tasks for the week.",
        "icon": "chart-no-axes-combined",
        "badge": "Most used",
        "default_name": "Weekly Team Summary",
    },
    {
        "type": Report.Type.TEAM_PERFORMANCE,
        "title": "Team Performance",
        "description": "Individual output, team momentum and productivity insights.",
        "icon": "users",
        "default_name": "Team Performance Review",
    },
    {
        "type": Report.Type.PROJECT_STATUS,
        "title": "Project Status",
        "description": "Project health, delivery risks and milestone tracking in one view.",
        "icon": "folder",
        "default_name": "Project Status Review",
    },
    {
        "type": Report.Type.TIME_TRACKING,
        "title": "Time Tracking",
        "description": "Task activity by team members across active projects.",
        "icon": "clock-3",
        "default_name": "Team Activity Report",
        "limitations": ["No time-entry model is configured; task activity is shown instead of logged hours."],
    },
    {
        "type": Report.Type.CUSTOM,
        "title": "Create custom report",
        "description": "Build a report tailored to your specific needs.",
        "icon": "circle-plus",
        "default_name": "Custom Report",
    },
)


def _date_range(parameters):
    today = timezone.localdate()
    end = parameters.get("end_date") or today
    if isinstance(end, str):
        end = date.fromisoformat(end)
    start = parameters.get("start_date") or (end - timedelta(days=6))
    if isinstance(start, str):
        start = date.fromisoformat(start)
    return start, end


def _tasks_for(report, user):
    parameters = report.parameters or {}
    tasks = visible_tasks_for(
        Task.objects.filter(department=report.department, is_archived=False)
        .select_related("project", "main_assignee")
        .prefetch_related("assignees"),
        user,
    )
    project_ids = parameters.get("projects") or []
    if project_ids:
        tasks = tasks.filter(project_id__in=project_ids)
    return tasks


def _weekly_result(report, user):
    start, end = _date_range(report.parameters)
    tasks = _tasks_for(report, user)
    completed = tasks.filter(status=Task.Status.COMPLETED, completed_at__date__range=(start, end))
    active = tasks.filter(status__in=(Task.Status.IN_PROGRESS, Task.Status.ON_HOLD))
    overdue = tasks.exclude(status=Task.Status.COMPLETED).filter(due_date__lt=timezone.localdate())
    return {
        "period": {"start": start.isoformat(), "end": end.isoformat()},
        "summary": {
            "total_tasks": tasks.count(),
            "completed_in_period": completed.count(),
            "in_progress": active.count(),
            "overdue": overdue.count(),
        },
        "attention_items": [
            {
                "id": str(task.id),
                "title": task.title,
                "project": task.project.name if task.project else None,
                "owner": task.main_assignee.get_full_name() or task.main_assignee.email if task.main_assignee else None,
                "due_date": task.due_date.isoformat() if task.due_date else None,
                "status": task.status,
            }
            for task in overdue.order_by("due_date")[:10]
        ],
    }


def _team_result(report, user):
    tasks = _tasks_for(report, user)
    members = report.department.users.filter(is_active=True).annotate(
        assigned=Count("assigned_tasks", filter=Q(assigned_tasks__in=tasks), distinct=True),
        completed=Count("assigned_tasks", filter=Q(assigned_tasks__in=tasks, assigned_tasks__status=Task.Status.COMPLETED), distinct=True),
        overdue=Count("assigned_tasks", filter=Q(assigned_tasks__in=tasks, assigned_tasks__due_date__lt=timezone.localdate()) & ~Q(assigned_tasks__status=Task.Status.COMPLETED), distinct=True),
    )
    rows = []
    for member in members.order_by("first_name", "last_name", "email"):
        rows.append({
            "id": str(member.id),
            "name": member.get_full_name() or member.email,
            "job_title": member.job_title,
            "assigned": member.assigned,
            "completed": member.completed,
            "overdue": member.overdue,
            "completion_rate": round(member.completed / member.assigned * 100, 1) if member.assigned else 0,
        })
    return {
        "summary": {
            "team_members": len(rows),
            "assigned_tasks": sum(row["assigned"] for row in rows),
            "completed_tasks": sum(row["completed"] for row in rows),
            "overdue_tasks": sum(row["overdue"] for row in rows),
        },
        "members": rows,
    }


def _project_result(report, user):
    tasks = _tasks_for(report, user)
    projects = Project.objects.filter(department=report.department).exclude(status=Project.Status.ARCHIVED)
    selected = report.parameters.get("projects") or []
    if selected:
        projects = projects.filter(id__in=selected)
    rows = []
    today = timezone.localdate()
    for project in projects.order_by("end_date", "name"):
        project_tasks = tasks.filter(project=project)
        total = project_tasks.count()
        completed = project_tasks.filter(status=Task.Status.COMPLETED).count()
        overdue = project_tasks.exclude(status=Task.Status.COMPLETED).filter(due_date__lt=today).count()
        progress = round(completed / total * 100) if total else 0
        health = "delayed" if project.end_date and project.end_date < today and project.status != Project.Status.COMPLETED else "at_risk" if overdue else "on_track"
        rows.append({
            "id": str(project.id), "name": project.name, "status": project.status,
            "health": health, "progress": progress, "total_tasks": total,
            "completed_tasks": completed, "overdue_tasks": overdue,
            "end_date": project.end_date.isoformat() if project.end_date else None,
            "manager": project.manager.get_full_name() or project.manager.email if project.manager else None,
        })
    return {
        "summary": {
            "active_projects": len(rows),
            "on_track": sum(row["health"] == "on_track" for row in rows),
            "at_risk": sum(row["health"] == "at_risk" for row in rows),
            "delayed": sum(row["health"] == "delayed" for row in rows),
        },
        "projects": rows,
    }


def _time_result(report, user):
    tasks = _tasks_for(report, user)
    by_member = defaultdict(lambda: {"assigned": 0, "completed": 0, "effort_points": 0})
    for task in tasks:
        for member in task.assignees.all():
            row = by_member[str(member.id)]
            row["id"] = str(member.id)
            row["name"] = member.get_full_name() or member.email
            row["assigned"] += 1
            row["effort_points"] += task.effort_score
            if task.status == Task.Status.COMPLETED:
                row["completed"] += 1
    return {
        "summary": {
            "tracked_hours": None,
            "active_tasks": tasks.exclude(status=Task.Status.COMPLETED).count(),
            "completed_tasks": tasks.filter(status=Task.Status.COMPLETED).count(),
            "effort_points": tasks.aggregate(total=Sum("effort_score"))["total"] or 0,
        },
        "members": list(by_member.values()),
        "meta": {"time_tracking_available": False, "reason": "Time entries are not present in the current data model."},
    }


def build_report_result(report, user):
    builders = {
        Report.Type.WEEKLY_PROGRESS: _weekly_result,
        Report.Type.TEAM_PERFORMANCE: _team_result,
        Report.Type.PROJECT_STATUS: _project_result,
        Report.Type.TIME_TRACKING: _time_result,
    }
    if report.report_type == Report.Type.CUSTOM:
        result = _weekly_result(report, user)
        result["custom_fields"] = report.parameters.get("fields", [])
        return result
    return builders[report.report_type](report, user)


def _add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = "" if value is None else str(value)


def build_report_docx(report, result):
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.8)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(report.name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(23, 50, 77)
    subtitle = document.add_paragraph(f"{report.get_report_type_display()} · Generated {timezone.localtime():%b %d, %Y %H:%M}")
    subtitle.runs[0].font.color.rgb = RGBColor(102, 112, 133)
    document.add_heading("Summary", level=1)
    _add_table(document, ("Metric", "Value"), [(key.replace("_", " ").title(), value) for key, value in result.get("summary", {}).items()])
    if result.get("attention_items"):
        document.add_heading("Tasks requiring attention", level=1)
        _add_table(document, ("Task", "Project", "Owner", "Due date", "Status"), [(x["title"], x["project"], x["owner"], x["due_date"], x["status"]) for x in result["attention_items"]])
    if result.get("members"):
        document.add_heading("Team member breakdown", level=1)
        _add_table(document, ("Team member", "Assigned", "Completed", "Overdue / Effort"), [(x["name"], x["assigned"], x["completed"], x.get("overdue", x.get("effort_points", 0))) for x in result["members"]])
    if result.get("projects"):
        document.add_heading("Project breakdown", level=1)
        _add_table(document, ("Project", "Health", "Progress", "Tasks", "Due date"), [(x["name"], x["health"].replace("_", " ").title(), f'{x["progress"]}%', x["total_tasks"], x["end_date"]) for x in result["projects"]])
    if result.get("meta", {}).get("time_tracking_available") is False:
        document.add_paragraph("Note: Logged-hour data is unavailable. This report uses task activity and effort as a proxy.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()
